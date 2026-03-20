"""
api.py  –  FastAPI backend for Job Market Intelligence (v2 dashboard)
─────────────────────────────────────────────────────────────────────
Run:   python -m uvicorn api:app --host 0.0.0.0 --port 8502
Open:  http://localhost:8502
"""
import os, sys, threading, json, re, io, time
from datetime import date, datetime
from typing import Optional

import numpy as np
import pandas as pd
import psycopg2
from fastapi import FastAPI, BackgroundTasks, Query, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from config.config import DB_CONFIG

app = FastAPI(title="Job Market Intelligence API", version="2.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ── DB helpers ────────────────────────────────────────────────────────────────

def _conn():
    return psycopg2.connect(**DB_CONFIG)

def _q(sql, params=None) -> pd.DataFrame:
    conn = _conn()
    try:
        return pd.read_sql_query(sql, conn, params=params)
    finally:
        conn.close()

def _safe(df: pd.DataFrame) -> list:
    """DataFrame → JSON-safe list.
    Uses pandas own JSON serializer to correctly handle numpy int64/float64/NaN/NaT.
    """
    return json.loads(df.to_json(orient="records", date_format="iso", default_handler=str))

# ── Skills extraction ─────────────────────────────────────────────────────────

SKILLS_VOCAB = [
    # Languages
    'Python','SQL','R','Scala','Java','JavaScript','Go','Bash','TypeScript','C++','C#',
    'Ruby','Rust','MATLAB','SAS','Swift','Kotlin','PHP','Perl',
    # Cloud & data platforms
    'AWS','Azure','GCP','Snowflake','Databricks','Redshift','BigQuery','S3','Lambda',
    'EMR','Glue','Athena','Synapse','Vertex AI','SageMaker','Dataflow','Pub/Sub',
    # ML / AI
    'TensorFlow','PyTorch','Scikit-learn','Keras','XGBoost','LightGBM','CatBoost',
    'Hugging Face','LangChain','OpenAI','LLM','Transformers','BERT','GPT','RAG',
    'Vector Database','Pinecone','Weaviate','MLflow','MLOps','ONNX','AutoML',
    # Orchestration / ETL
    'Airflow','dbt','Spark','Kafka','Luigi','Prefect','Dagster','NiFi',
    'Informatica','Talend','Fivetran','Stitch','Mage','ETL','ELT','Flink',
    # Databases
    'PostgreSQL','MySQL','MongoDB','Elasticsearch','Redis','Cassandra','DynamoDB',
    'SQLite','Oracle','SQL Server','MariaDB','Neo4j','Cosmos DB','Hive','Presto','Trino',
    # Visualization
    'Tableau','Power BI','Looker','Plotly','Matplotlib','Seaborn','Grafana',
    'Superset','Qlik','D3.js','Streamlit','Dash',
    # DevOps / Infra
    'Docker','Kubernetes','Terraform','Git','CI/CD','Jenkins','GitHub Actions',
    'Ansible','Helm','Linux','Prometheus','Datadog',
    # Frameworks / Libraries
    'Pandas','NumPy','Flask','FastAPI','Django','React','Node.js','Spring',
    'Hadoop','Spark','Kafka',
    # Concepts / Methods
    'Machine Learning','Deep Learning','NLP','Computer Vision','Reinforcement Learning',
    'Data Engineering','Data Science','Data Analysis','Statistics','Probability',
    'A/B Testing','Regression','Classification','Clustering','Neural Network',
    'Data Warehouse','Data Lake','Data Lakehouse','Feature Engineering',
    'Model Deployment','REST API','GraphQL','Agile','Scrum','Excel',
    'Data Modeling','Data Governance','Data Quality','Prompt Engineering',
    'Time Series','Forecasting','Recommendation System',
]

# Build lowercase → display lookup once
_SKILL_LOWER_MAP = {s.lower(): s for s in SKILLS_VOCAB}

def _extract_skills_from_text(text: str) -> list:
    if not text:
        return []
    text_lower = text.lower()
    found = []
    for skill_lower, skill_display in _SKILL_LOWER_MAP.items():
        pattern = r'\b' + re.escape(skill_lower) + r'\b'
        if re.search(pattern, text_lower):
            found.append(skill_display)
    return found

def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF resume bytes.
    Priority: PyMuPDF (best for resumes) → pdfminer → pypdf → raw bytes.
    Tries multiple PyMuPDF modes to maximise extraction from all PDF types.
    """
    # ── 1. PyMuPDF — handles ATS-style, multi-column, modern PDFs ────────────
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        # Try "text" mode first — best for most resumes
        parts_text = []
        for page in doc:
            parts_text.append(page.get_text("text"))
        text_mode = "\n".join(parts_text).strip()

        if text_mode:
            doc.close()
            return text_mode

        # Fallback: "blocks" mode — different layout algorithm, catches more
        parts_blocks = []
        for page in doc:
            for block in page.get_text("blocks"):
                if block[6] == 0:  # text block (not image)
                    parts_blocks.append(block[4])
        text_blocks = "\n".join(parts_blocks).strip()

        doc.close()
        if text_blocks:
            return text_blocks

        print("[pdf] PyMuPDF: no text found — PDF may be image-based (scanned)")

    except Exception as e:
        print(f"[pdf] PyMuPDF failed: {e}")

    # ── 2. pdfminer.six ───────────────────────────────────────────────────────
    try:
        from pdfminer.high_level import extract_text as _pdfminer
        text = _pdfminer(io.BytesIO(content))
        if text and text.strip():
            return text
    except Exception:
        pass

    # ── 3. pypdf ──────────────────────────────────────────────────────────────
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text
    except Exception:
        pass

    # ── 4. Last resort: decode as UTF-8 ──────────────────────────────────────
    return content.decode("utf-8", errors="ignore")


def _extract_docx_text(content: bytes) -> str:
    """Extract text from a .docx file.
    Primary:   raw XML <w:t> tag parsing  (works on ALL docx variants, no library needed)
    Fallback:  python-docx paragraph/table walk
    """
    import zipfile, re as _re

    # ── 1. Raw ZIP/XML extraction — most reliable across all DOCX variants ─────
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            names = z.namelist()
            parts = []
            # Process in order: main doc, headers, footers, footnotes
            targets = ['word/document.xml', 'word/header1.xml', 'word/footer1.xml']
            for target in targets:
                if target in names:
                    xml = z.read(target).decode('utf-8', errors='ignore')
                    # Extract text from <w:t> runs — the authoritative text nodes in OOXML
                    texts = _re.findall(r'<w:t(?:\s[^>]*)?>([^<]*)</w:t>', xml)
                    for t in texts:
                        t = t.strip()
                        if t:
                            parts.append(t)
            if parts:
                # Join with spaces; paragraph breaks come naturally from structure
                text = ' '.join(parts)
                print(f"[docx-xml] extracted {len(text)} chars via raw XML")
                return text
    except Exception as e:
        print(f"[docx-xml] failed: {e}")

    # ── 2. python-docx paragraph walk ──────────────────────────────────────────
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text)
        text = "\n".join(lines)
        if text.strip():
            print(f"[docx-lib] extracted {len(text)} chars via python-docx")
            return text
    except Exception as e:
        print(f"[docx-lib] python-docx failed: {e}")

    return ""  # empty — caller will report the error properly

# Extraction progress tracker
_extract_progress = {"running": False, "total": 0, "done": 0, "msg": ""}

# ── LLM helpers (OpenAI) ──────────────────────────────────────────────────────

def _llm_analyze(resume_text: str, job_title: str, job_company: str,
                 job_desc: str, job_skills_list: list, api_key: str) -> dict:
    """Use GPT-4o-mini to compare resume vs job. Returns structured dict."""
    try:
        import openai
        client = openai.OpenAI(api_key=api_key)
        skills_str = ", ".join(job_skills_list[:30]) if job_skills_list else "not specified"
        prompt = f"""You are a senior technical recruiter. Analyze this resume against the job posting.

JOB POSTING:
Title: {job_title}
Company: {job_company}
Required Skills: {skills_str}
Description: {(job_desc or 'No description available')[:3000]}

CANDIDATE RESUME:
{resume_text[:3000]}

Return ONLY a valid JSON object with these exact keys:
{{
  "match_score": <integer 0-100, be realistic and precise>,
  "verdict": "<one of: Strongly Recommended | Recommended | Consider Applying | Not Recommended>",
  "summary": "<2-3 sentences: direct, specific overall assessment of fit>",
  "strengths": ["<key strength relevant to this job>", ...],
  "gaps": ["<specific missing qualification or skill>", ...],
  "matched_skills": ["<skill found in both resume and job>", ...],
  "missing_skills": ["<skill required by job but not in resume>", ...],
  "suggestions": ["<specific, actionable improvement>", ...],
  "interview_tips": ["<specific topic or question to prepare for based on this job>", ...]
}}"""
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.2,
            max_tokens=1600,
        )
        result = json.loads(resp.choices[0].message.content)
        result["llm_used"] = True
        return result
    except Exception as e:
        return {"error": str(e), "llm_used": False}


def _llm_cover_letter(resume_text: str, job_title: str, job_company: str,
                      job_desc: str, api_key: str) -> str:
    """Generate a tailored cover letter using GPT."""
    try:
        import openai
        client = openai.OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": f"""Write a professional, concise cover letter (3 short paragraphs) for:
Job: {job_title} at {job_company}
Job Description: {(job_desc or '')[:1500]}
Candidate Resume: {resume_text[:1500]}

Rules: Be specific and highlight genuine matches. No generic filler. No placeholders.
Start with "Dear Hiring Manager," — end with a confident call to action."""}],
            temperature=0.4,
            max_tokens=700,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error generating cover letter: {e}"


# ── Scrape config helpers ─────────────────────────────────────────────────────

CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scrape_config.json")

# High-density tech job markets — scrape these first
DEFAULT_PRIORITY_LOCATIONS = [
    "San Francisco, CA", "Seattle, WA", "New York, NY", "Austin, TX",
    "Los Angeles, CA", "Chicago, IL", "Boston, MA", "Denver, CO",
    "Atlanta, GA", "Washington, DC", "Dallas, TX", "San Diego, CA",
    "Charlotte, NC", "Raleigh, NC", "Minneapolis, MN", "Miami, FL",
    "Phoenix, AZ", "Nashville, TN", "Philadelphia, PA", "Remote",
]

def _load_scrape_config():
    try:
        from config.config import SEARCH_QUERIES
        default_roles = SEARCH_QUERIES
    except Exception:
        default_roles = ["Data Engineer", "Data Analyst", "Data Scientist",
                         "Machine Learning Engineer", "AI Engineer"]
    default = {
        "roles": default_roles,
        "priority_locations": DEFAULT_PRIORITY_LOCATIONS,
        "schedule_enabled": False,
        "schedule_interval_hours": 24,
        "openai_api_key": "",
    }
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH) as f:
                stored = json.load(f)
                default = {**default, **stored}
        except Exception:
            pass
    # Allow OPENAI_API_KEY env var to override (used on Railway / any cloud host)
    env_key = os.getenv("OPENAI_API_KEY", "").strip()
    if env_key:
        default["openai_api_key"] = env_key
    return default

# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/api/stats")
def stats():
    df = _q("""
        SELECT
            COUNT(*)                                                          AS total_jobs,
            COUNT(DISTINCT company)                                           AS total_companies,
            ROUND(AVG((salary_min+salary_max)/2))                             AS avg_salary,
            ROUND(100.0*SUM(is_remote::int)/NULLIF(COUNT(*),0),1)            AS pct_remote,
            TO_CHAR(MAX(date_scraped)::date, 'YYYY-MM-DD')                   AS last_scraped,
            COUNT(CASE WHEN date_scraped >= NOW()-INTERVAL '24 hours' THEN 1 END) AS scraped_today
        FROM jobs WHERE is_active=TRUE
    """)
    if df.empty:
        return {}
    return _safe(df)[0]

@app.get("/api/jobs-over-time")
def jobs_over_time():
    df = _q("""
        SELECT date_scraped::date AS day, COUNT(*) AS count
        FROM jobs WHERE date_scraped >= NOW()-INTERVAL '30 days'
        GROUP BY day ORDER BY day
    """)
    df["day"] = df["day"].astype(str)
    df["count"] = df["count"].astype(int)
    return df.to_dict("records")

@app.get("/api/top-roles")
def top_roles(n: int = 8):
    df = _q(f"""
        SELECT search_query AS role, COUNT(*) AS count,
               COALESCE(ROUND(AVG((salary_min+salary_max)/2)),0) AS avg_salary
        FROM jobs WHERE is_active=TRUE
        GROUP BY search_query ORDER BY count DESC LIMIT {n}
    """)
    return _safe(df)

@app.get("/api/top-companies")
def top_companies(n: int = 15):
    df = _q(f"""
        SELECT company, COUNT(*) AS job_count,
               ROUND(AVG((salary_min+salary_max)/2)) AS avg_salary
        FROM jobs
        WHERE is_active=TRUE AND company IS NOT NULL AND company NOT IN ('Unknown','')
        GROUP BY company ORDER BY job_count DESC LIMIT {n}
    """)
    return _safe(df)

@app.get("/api/jobs-by-state")
def jobs_by_state():
    df = _q("""
        SELECT TRIM(SPLIT_PART(location,',',2)) AS state_raw,
               COUNT(*) AS job_count,
               ROUND(AVG((salary_min+salary_max)/2)) AS avg_salary
        FROM jobs
        WHERE is_active=TRUE AND location IS NOT NULL
          AND location NOT ILIKE '%remote%' AND location LIKE '%,%'
        GROUP BY state_raw HAVING COUNT(*) > 1
        ORDER BY job_count DESC
    """)
    return _safe(df)

@app.get("/api/salary-by-role")
def salary_by_role():
    df = _q("""
        SELECT search_query AS role,
               salary_min, salary_max,
               (salary_min+salary_max)/2 AS salary_mid,
               seniority, company
        FROM jobs
        WHERE is_active=TRUE
          AND salary_min IS NOT NULL AND salary_max IS NOT NULL
          AND salary_max BETWEEN 20000 AND 500000
    """)
    return _safe(df)

@app.get("/api/salary-by-seniority")
def salary_by_seniority():
    df = _q("""
        SELECT seniority,
               ROUND(AVG((salary_min+salary_max)/2)) AS avg_salary,
               ROUND(MIN(salary_min))                AS min_salary,
               ROUND(MAX(salary_max))                AS max_salary,
               COUNT(*)                              AS count
        FROM jobs
        WHERE is_active=TRUE AND salary_min IS NOT NULL
          AND seniority IS NOT NULL AND seniority != ''
        GROUP BY seniority ORDER BY avg_salary DESC
    """)
    return _safe(df)

@app.get("/api/top-skills")
def top_skills(n: int = 30):
    df = _q(f"SELECT skill, COUNT(*) AS count FROM job_skills GROUP BY skill ORDER BY count DESC LIMIT {n}")
    return _safe(df)

@app.get("/api/jobs")
def jobs(
    role:        Optional[str]   = None,
    location:    Optional[str]   = None,
    source:      Optional[str]   = None,
    remote_only: bool            = False,
    min_sal:     Optional[float] = None,
    max_sal:     Optional[float] = None,
    days_back:   int             = 30,
    hours_back:  Optional[int]   = None,
    search:      Optional[str]   = None,
    limit:       int             = 50,
    offset:      int             = 0,
):
    conds  = ["is_active=TRUE"]
    params = []
    # hours_back=0 → all time (no date filter); >0 → filter; None → use days_back
    if hours_back is None:
        conds.append(f"date_scraped >= NOW()-INTERVAL '{days_back} days'")
    elif hours_back > 0:
        conds.append(f"date_scraped >= NOW()-INTERVAL '{hours_back} hours'")
    # hours_back == 0 → no date filter (all time)
    if role:        conds.append("search_query=%s");             params.append(role)
    if location:    conds.append("location ILIKE %s");           params.append(f"%{location}%")
    if source:      conds.append("source=%s");                   params.append(source)
    if remote_only: conds.append("is_remote=TRUE")
    if min_sal:     conds.append("salary_min>=%s");              params.append(min_sal)
    if max_sal:     conds.append("salary_max<=%s");              params.append(max_sal)
    if search:
        conds.append("(title ILIKE %s OR company ILIKE %s)")
        params += [f"%{search}%", f"%{search}%"]
    df = _q(
        f"SELECT job_id,title,company,location,is_remote,salary_min,salary_max,"
        f"seniority,source,date_posted,url,search_query,date_scraped "
        f"FROM jobs WHERE {' AND '.join(conds)} "
        f"ORDER BY date_scraped DESC LIMIT {limit} OFFSET {offset}",
        params or None
    )
    return _safe(df)

@app.get("/api/jobs/count")
def jobs_count(
    role:        Optional[str]   = None,
    location:    Optional[str]   = None,
    source:      Optional[str]   = None,
    remote_only: bool            = False,
    min_sal:     Optional[float] = None,
    max_sal:     Optional[float] = None,
    days_back:   int             = 30,
    hours_back:  Optional[int]   = None,
    search:      Optional[str]   = None,
):
    conds  = ["is_active=TRUE"]
    params = []
    if hours_back is None:
        conds.append(f"date_scraped >= NOW()-INTERVAL '{days_back} days'")
    elif hours_back > 0:
        conds.append(f"date_scraped >= NOW()-INTERVAL '{hours_back} hours'")
    # hours_back == 0 → no date filter (all time)
    if role:        conds.append("search_query=%s");  params.append(role)
    if location:    conds.append("location ILIKE %s"); params.append(f"%{location}%")
    if source:      conds.append("source=%s");         params.append(source)
    if remote_only: conds.append("is_remote=TRUE")
    if min_sal:     conds.append("salary_min>=%s");   params.append(min_sal)
    if max_sal:     conds.append("salary_max<=%s");   params.append(max_sal)
    if search:
        conds.append("(title ILIKE %s OR company ILIKE %s)")
        params += [f"%{search}%", f"%{search}%"]
    df = _q(f"SELECT COUNT(*) AS cnt FROM jobs WHERE {' AND '.join(conds)}", params or None)
    return {"count": int(df["cnt"].iloc[0]) if not df.empty else 0}

@app.get("/api/job/{job_id}")
def get_job(job_id: str):
    """Get full job details including description."""
    df = _q(
        "SELECT job_id,title,company,location,is_remote,salary_min,salary_max,"
        "seniority,source,date_posted,url,search_query,date_scraped,description "
        "FROM jobs WHERE job_id=%s",
        [job_id]
    )
    if df.empty:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return _safe(df)[0]

@app.get("/api/scrape-history")
def scrape_history(n: int = 25):
    df = _q(f"""
        SELECT run_timestamp::date AS date, run_timestamp::time(0) AS time,
               source, query, location, jobs_found, jobs_new, status, error_msg
        FROM scrape_runs ORDER BY run_timestamp DESC LIMIT {n}
    """)
    return _safe(df)

@app.get("/api/roles")
def roles():
    df = _q("SELECT DISTINCT search_query FROM jobs WHERE is_active=TRUE ORDER BY search_query")
    return df["search_query"].tolist() if not df.empty else []

@app.get("/api/sources")
def sources():
    df = _q("SELECT DISTINCT source FROM jobs WHERE is_active=TRUE ORDER BY source")
    return df["source"].tolist() if not df.empty else []

@app.get("/api/db-stats")
def db_stats():
    df = _q("""
        SELECT (SELECT COUNT(*) FROM jobs)        AS total_jobs,
               (SELECT COUNT(*) FROM job_skills)  AS total_skills,
               (SELECT COUNT(*) FROM scrape_runs) AS total_runs,
               (SELECT COUNT(*) FROM scrape_runs WHERE status='failed') AS failed_runs
    """)
    return _safe(df)[0] if not df.empty else {}

# ── Skills extraction endpoints ───────────────────────────────────────────────

@app.post("/api/skills/extract")
def skills_extract(bt: BackgroundTasks):
    """Background skill extraction from all job descriptions."""
    if _extract_progress["running"]:
        return {"status": "already_running", "message": "Extraction already in progress"}

    def _go():
        _extract_progress["running"] = True
        _extract_progress["done"]    = 0
        _extract_progress["total"]   = 0
        _extract_progress["msg"]     = "Starting skill extraction..."
        try:
            conn = _conn()
            cur  = conn.cursor()

            # Determine if description column exists
            cur.execute("""
                SELECT column_name FROM information_schema.columns
                WHERE table_name='jobs' AND column_name='description'
            """)
            has_desc = cur.fetchone() is not None

            if has_desc:
                cur.execute(
                    "SELECT job_id, description FROM jobs "
                    "WHERE is_active=TRUE AND description IS NOT NULL AND description != ''"
                )
            else:
                cur.execute(
                    "SELECT job_id, title || ' ' || COALESCE(search_query,'') AS description "
                    "FROM jobs WHERE is_active=TRUE"
                )
            rows = cur.fetchall()
            _extract_progress["total"] = len(rows)
            _extract_progress["msg"]   = f"Processing {len(rows)} jobs…"

            # Clear existing extracted skills
            cur.execute("DELETE FROM job_skills")
            conn.commit()

            batch = []
            for i, (job_id, desc) in enumerate(rows):
                skills = _extract_skills_from_text(desc or "")
                for skill in skills:
                    batch.append((job_id, skill))
                _extract_progress["done"] = i + 1
                # Batch insert every 200 rows
                if len(batch) >= 200:
                    cur.executemany(
                        "INSERT INTO job_skills (job_id, skill) VALUES (%s,%s) ON CONFLICT DO NOTHING",
                        batch
                    )
                    conn.commit()
                    batch = []

            if batch:
                cur.executemany(
                    "INSERT INTO job_skills (job_id, skill) VALUES (%s,%s) ON CONFLICT DO NOTHING",
                    batch
                )
                conn.commit()

            conn.close()
            _extract_progress["msg"] = (
                f"✅ Done! Extracted skills from {len(rows)} jobs."
            )
        except Exception as e:
            _extract_progress["msg"] = f"❌ Error: {e}"
        finally:
            _extract_progress["running"] = False

    bt.add_task(_go)
    return {"status": "started", "message": "Skill extraction started in background"}

@app.get("/api/skills/extract/status")
def skills_extract_status():
    return _extract_progress

# ── Scrape config endpoints ───────────────────────────────────────────────────

@app.get("/api/config")
def get_config():
    cfg = _load_scrape_config()
    # Never expose the raw key — only tell the UI whether one is set
    key = cfg.get("openai_api_key", "")
    cfg["openai_key_set"]  = bool(key)
    cfg["openai_api_key"]  = "***hidden***" if key else ""
    return cfg

@app.post("/api/config")
async def save_config(request: Request):
    body = await request.json()
    # Don't overwrite the real key if the masked placeholder is sent
    if body.get("openai_api_key") == "***hidden***":
        existing = _load_scrape_config()
        body["openai_api_key"] = existing.get("openai_api_key", "")
    with open(CONFIG_PATH, "w") as f:
        json.dump(body, f, indent=2)
    return {"status": "saved"}

@app.post("/api/config/openai-key")
async def save_openai_key(request: Request):
    """Dedicated endpoint to save the OpenAI API key securely."""
    body = await request.json()
    key  = body.get("key", "").strip()
    cfg  = _load_scrape_config()
    cfg["openai_api_key"] = key
    with open(CONFIG_PATH, "w") as f:
        json.dump(cfg, f, indent=2)
    return {"status": "saved", "key_set": bool(key)}

# ── Resume analysis endpoint ──────────────────────────────────────────────────

def _is_human_readable(text: str) -> bool:
    """Return True if text looks like real words rather than binary garbage."""
    if not text or len(text.strip()) < 30:
        return False
    sample = text[:600]
    # Count printable ASCII + common unicode letters vs raw binary
    printable = sum(1 for c in sample if c.isprintable() or c in '\n\r\t')
    ratio = printable / len(sample)
    return ratio >= 0.85   # 85%+ printable characters = real text


@app.post("/api/resume/test-extract")
async def resume_test_extract(resume: UploadFile = File(...)):
    """Debug endpoint — returns raw extracted text so you can verify file reading."""
    content = await resume.read()
    fname   = (resume.filename or "").lower()
    ctype   = (resume.content_type or "").lower()
    is_pdf  = fname.endswith(".pdf")  or "pdf" in ctype
    is_docx = fname.endswith(".docx") or "wordprocessingml" in ctype or "openxmlformats" in ctype
    is_doc  = fname.endswith(".doc")  and not is_docx

    if is_pdf:
        text = _extract_pdf_text(content)
        method = "PyMuPDF + pdfminer fallback"
    elif is_docx or is_doc:
        text = _extract_docx_text(content)
        method = "DOCX XML parser + python-docx fallback"
    else:
        text = content.decode("utf-8", errors="ignore")
        method = "plain-text decode"

    readable = _is_human_readable(text)
    skills   = _extract_skills_from_text(text) if readable else []

    return {
        "filename":        resume.filename,
        "content_type":    resume.content_type,
        "file_size_bytes": len(content),
        "file_type":       "PDF" if is_pdf else ("DOCX/DOC" if (is_docx or is_doc) else "Text"),
        "method":          method,
        "chars_extracted": len(text.strip()),
        "skills_found":    skills,
        "preview":         text[:1500].strip() if readable else "(binary/unreadable content — extraction failed)",
        "readable":        readable,
    }


@app.post("/api/resume/analyze")
async def resume_analyze(
    resume:                UploadFile    = File(...),
    job_id:                Optional[str] = Form(None),
    generate_cover_letter: bool          = Form(False),
):
    """
    Analyze resume vs a scraped job posting.
    Uses GPT-4o-mini when an OpenAI API key is configured, otherwise falls
    back to regex-based skill matching.
    """
    content = await resume.read()

    # ── Extract resume text ───────────────────────────────────────────────────
    fname = (resume.filename or "").lower()
    ctype = (resume.content_type or "").lower()

    is_pdf  = fname.endswith(".pdf")  or "pdf"  in ctype
    is_docx = fname.endswith(".docx") or "wordprocessingml" in ctype or "openxmlformats" in ctype
    is_doc  = fname.endswith(".doc")  and not is_docx

    if is_pdf:
        resume_text = _extract_pdf_text(content)
    elif is_docx:
        resume_text = _extract_docx_text(content)
    elif is_doc:
        # Old .doc format — try docx extractor anyway, fall back to plain text
        resume_text = _extract_docx_text(content)
        if not resume_text.strip():
            resume_text = content.decode("utf-8", errors="ignore")
    else:
        # Plain text / unknown — decode directly
        resume_text = content.decode("utf-8", errors="ignore")

    print(f"[resume] file={fname!r} size={len(content)}B extracted={len(resume_text.strip())} chars")

    if not resume_text.strip():
        return JSONResponse({
            "error": "Could not extract text from your resume. "
                     "If it's a scanned/image-based PDF, please save it as plain text first. "
                     "A Word (.docx) or plain text (.txt) file will always work."
        }, status_code=400)

    resume_skills = set(_extract_skills_from_text(resume_text))

    # ── Fetch job info ────────────────────────────────────────────────────────
    job_info    = {}
    job_desc    = ""
    job_skills  = set()

    if job_id:
        try:
            df = _q("SELECT title, company, location, description "
                    "FROM jobs WHERE job_id=%s", [job_id])
            if not df.empty:
                row = df.iloc[0]
                job_info = {
                    "title":    str(row.get("title",    "") or ""),
                    "company":  str(row.get("company",  "") or ""),
                    "location": str(row.get("location", "") or ""),
                }
                job_desc = str(row.get("description") or "")
                if job_desc:
                    job_skills.update(_extract_skills_from_text(job_desc))
        except Exception:
            pass
        try:
            df2 = _q("SELECT skill FROM job_skills WHERE job_id=%s", [job_id])
            if not df2.empty:
                job_skills.update(df2["skill"].tolist())
        except Exception:
            pass

    # ── Try LLM analysis first ────────────────────────────────────────────────
    cfg     = _load_scrape_config()
    api_key = cfg.get("openai_api_key", "").strip()
    llm_result = {}

    if api_key:
        llm_result = _llm_analyze(
            resume_text   = resume_text,
            job_title     = job_info.get("title",   "Unknown Role"),
            job_company   = job_info.get("company", "Unknown Company"),
            job_desc      = job_desc,
            job_skills_list = sorted(job_skills),
            api_key       = api_key,
        )

    # ── Regex fallback (always computed, used as backup) ──────────────────────
    matched = sorted(resume_skills & job_skills)
    missing = sorted(job_skills - resume_skills)
    extra   = sorted(resume_skills - job_skills)
    regex_score = round(len(matched) / len(job_skills) * 100) if job_skills else 0

    # ── Merge results ─────────────────────────────────────────────────────────
    if llm_result.get("llm_used"):
        # LLM provides richer data — override regex for key fields
        final = {
            "llm_used":       True,
            "match_score":    llm_result.get("match_score", regex_score),
            "verdict":        llm_result.get("verdict", ""),
            "summary":        llm_result.get("summary", ""),
            "strengths":      llm_result.get("strengths", []),
            "gaps":           llm_result.get("gaps", []),
            "matched_skills": llm_result.get("matched_skills", matched),
            "missing_skills": llm_result.get("missing_skills", missing),
            "extra_skills":   extra,
            "suggestions":    llm_result.get("suggestions", []),
            "interview_tips": llm_result.get("interview_tips", []),
            "job_info":       job_info,
            "resume_skills":  sorted(resume_skills),
            "job_skills":     sorted(job_skills),
        }
    else:
        # Regex-only fallback
        suggestions = []
        if missing:
            suggestions.append(f"Focus on learning: {', '.join(missing[:5])}.")
        label = ("Strong match — well qualified!" if regex_score >= 80
                 else "Good match — a few skill gaps to bridge." if regex_score >= 60
                 else "Moderate match — significant upskilling recommended." if regex_score >= 40
                 else "Low match — consider roles that better align with your skills.")
        suggestions.append(label)
        final = {
            "llm_used":       False,
            "match_score":    regex_score,
            "verdict":        "",
            "summary":        label,
            "strengths":      [],
            "gaps":           [],
            "matched_skills": matched,
            "missing_skills": missing,
            "extra_skills":   extra,
            "suggestions":    suggestions,
            "interview_tips": [],
            "job_info":       job_info,
            "resume_skills":  sorted(resume_skills),
            "job_skills":     sorted(job_skills),
        }
        if not api_key:
            final["llm_note"] = "Add your OpenAI API key in Pipeline → Settings to unlock AI-powered analysis."

    # ── Optional cover letter ─────────────────────────────────────────────────
    if generate_cover_letter and api_key:
        final["cover_letter"] = _llm_cover_letter(
            resume_text = resume_text,
            job_title   = job_info.get("title",   "Unknown Role"),
            job_company = job_info.get("company", "Unknown Company"),
            job_desc    = job_desc,
            api_key     = api_key,
        )

    # Include truncated resume text for chat context
    final["resume_text_preview"] = resume_text[:2500]

    return final

# ── Chat endpoint ────────────────────────────────────────────────────────────

@app.post("/api/resume/chat")
async def resume_chat(request: Request):
    """Multi-turn chat about resume vs job fit. Maintains full message history."""
    body       = await request.json()
    messages   = body.get("messages", [])       # list of {role, content}
    context    = body.get("context", {})         # {resume_text, job_info, analysis_summary}

    cfg     = _load_scrape_config()
    api_key = cfg.get("openai_api_key", "").strip()
    if not api_key:
        return JSONResponse({"error": "No OpenAI API key configured. Add it in Pipeline → Settings."}, status_code=400)

    job    = context.get("job_info",         {})
    resume = context.get("resume_text",      "")[:3000]
    summ   = context.get("analysis_summary", "")

    system_prompt = f"""You are an expert career coach and technical recruiter.
The candidate is evaluating this job opportunity:

Job: {job.get('title','Unknown')} at {job.get('company','Unknown')}{(' — ' + job.get('location','')) if job.get('location') else ''}

Initial Analysis Summary:
{summ}

Resume excerpt:
{resume}

Help the candidate with: fit assessment, skill gaps, how to position themselves, interview prep,
salary negotiation strategy, or any career question about this opportunity.
Be direct, specific, and actionable. Reference specific details from the resume and job when relevant."""

    import openai
    client = openai.OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model    = "gpt-4o-mini",
        messages = [{"role": "system", "content": system_prompt}] + messages,
        temperature = 0.7,
        max_tokens  = 800,
    )
    return {"reply": resp.choices[0].message.content}


# ── Application Tracker ───────────────────────────────────────────────────────

def _init_applications_table():
    conn = _conn()
    cur  = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS job_applications (
            id           SERIAL PRIMARY KEY,
            job_id       VARCHAR(255),
            status       VARCHAR(50) DEFAULT 'saved',
            notes        TEXT        DEFAULT '',
            applied_date DATE,
            created_at   TIMESTAMP   DEFAULT NOW(),
            updated_at   TIMESTAMP   DEFAULT NOW()
        )
    """)
    conn.commit()
    conn.close()

try:
    _init_applications_table()
except Exception as _e:
    print(f"[startup] applications table: {_e}")

APP_STATUSES = ['saved', 'applied', 'phone_screen', 'technical', 'final_round', 'offer', 'rejected']

@app.get("/api/applications")
def get_applications():
    df = _q("""
        SELECT a.id, a.job_id, a.status, a.notes,
               TO_CHAR(a.applied_date, 'YYYY-MM-DD') AS applied_date,
               TO_CHAR(a.created_at,   'YYYY-MM-DD') AS created_at,
               TO_CHAR(a.updated_at,   'YYYY-MM-DD') AS updated_at,
               j.title, j.company, j.location,
               j.salary_min, j.salary_max, j.source, j.url, j.search_query, j.is_remote
        FROM job_applications a
        LEFT JOIN jobs j ON a.job_id = j.job_id
        ORDER BY a.updated_at DESC
    """)
    return _safe(df)

@app.post("/api/applications")
async def create_application(request: Request):
    body = await request.json()
    job_id = body.get("job_id")
    # Prevent duplicate saves
    existing = _q("SELECT id FROM job_applications WHERE job_id=%s", [job_id])
    if not existing.empty:
        return {"id": int(existing.iloc[0]["id"]), "status": "already_exists"}
    conn = _conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO job_applications (job_id, status, notes) VALUES (%s,%s,%s) RETURNING id",
        [job_id, body.get("status", "saved"), body.get("notes", "")]
    )
    new_id = cur.fetchone()[0]
    conn.commit(); conn.close()
    return {"id": new_id, "status": "created"}

@app.put("/api/applications/{app_id}")
async def update_application(app_id: int, request: Request):
    body = await request.json()
    sets, vals = [], []
    if "status" in body: sets.append("status=%s"); vals.append(body["status"])
    if "notes"  in body: sets.append("notes=%s");  vals.append(body["notes"])
    if "applied_date" in body: sets.append("applied_date=%s"); vals.append(body["applied_date"] or None)
    if not sets:
        return {"status": "nothing_to_update"}
    sets.append("updated_at=NOW()")
    vals.append(app_id)
    conn = _conn(); cur = conn.cursor()
    cur.execute(f"UPDATE job_applications SET {', '.join(sets)} WHERE id=%s", vals)
    conn.commit(); conn.close()
    return {"status": "updated"}

@app.delete("/api/applications/{app_id}")
def delete_application(app_id: int):
    conn = _conn(); cur = conn.cursor()
    cur.execute("DELETE FROM job_applications WHERE id=%s", [app_id])
    conn.commit(); conn.close()
    return {"status": "deleted"}

@app.get("/api/applications/check/{job_id}")
def check_application(job_id: str):
    df = _q("SELECT id, status FROM job_applications WHERE job_id=%s", [job_id])
    if df.empty:
        return {"saved": False}
    return {"saved": True, "id": int(df.iloc[0]["id"]), "status": str(df.iloc[0]["status"])}


# ── Batch Resume Scoring ──────────────────────────────────────────────────────

@app.post("/api/resume/batch-score")
async def batch_score(resume: UploadFile = File(...), limit: int = Form(50)):
    """Score resume against top N jobs by regex skill match. Fast, no LLM cost."""
    content = await resume.read()
    fname   = (resume.filename or "").lower()
    ctype   = (resume.content_type or "").lower()
    is_pdf  = fname.endswith(".pdf") or "pdf" in ctype
    is_docx = fname.endswith(".docx") or "wordprocessingml" in ctype
    if is_pdf:
        text = _extract_pdf_text(content)
    elif is_docx:
        text = _extract_docx_text(content)
    else:
        text = content.decode("utf-8", errors="ignore")
    if not text.strip():
        return JSONResponse({"error": "Could not extract resume text"}, status_code=400)

    resume_skills = set(_extract_skills_from_text(text))
    if not resume_skills:
        return JSONResponse({"error": "No recognizable skills found in resume"}, status_code=400)

    df = _q(f"""
        SELECT j.job_id, j.title, j.company, j.location,
               j.salary_min, j.salary_max, j.source, j.url, j.search_query, j.is_remote,
               STRING_AGG(js.skill, '|') AS skills_str
        FROM jobs j
        LEFT JOIN job_skills js ON j.job_id = js.job_id
        WHERE j.is_active=TRUE
        GROUP BY j.job_id, j.title, j.company, j.location,
                 j.salary_min, j.salary_max, j.source, j.url, j.search_query, j.is_remote
        ORDER BY j.date_scraped DESC
        LIMIT {int(limit) * 4}
    """)

    results = []
    for _, row in df.iterrows():
        skills_raw = row.get("skills_str") or ""
        job_skills = {s.strip() for s in skills_raw.split("|") if s.strip()}
        if not job_skills:
            continue
        matched = resume_skills & job_skills
        score   = round(len(matched) / len(job_skills) * 100)
        results.append({
            "job_id":        str(row["job_id"] or ""),
            "title":         str(row["title"]   or ""),
            "company":       str(row["company"] or ""),
            "location":      str(row["location"] or ""),
            "salary_min":    row["salary_min"],
            "salary_max":    row["salary_max"],
            "source":        str(row["source"]  or ""),
            "url":           str(row["url"]     or ""),
            "search_query":  str(row["search_query"] or ""),
            "is_remote":     bool(row["is_remote"]),
            "match_score":   score,
            "matched_count": len(matched),
            "total_skills":  len(job_skills),
            "matched_skills": sorted(matched),
        })

    results.sort(key=lambda x: x["match_score"], reverse=True)
    return json.loads(json.dumps(results[:int(limit)], default=str))


# ── Market Trends ─────────────────────────────────────────────────────────────

@app.get("/api/trends")
def market_trends(days: int = 60):
    df = _q(f"""
        SELECT date_scraped::date AS day,
               search_query       AS role,
               COUNT(*)           AS count
        FROM jobs
        WHERE date_scraped >= NOW()-INTERVAL '{days} days'
          AND search_query IS NOT NULL
        GROUP BY day, role
        ORDER BY day, role
    """)
    df["day"] = df["day"].astype(str)
    return _safe(df)


# ── Deduplication ──────────────────────────────────────────────────────────────

@app.get("/api/jobs/duplicate-count")
def duplicate_count():
    df = _q("""
        SELECT COUNT(*) AS cnt FROM (
            SELECT j1.job_id
            FROM jobs j1
            JOIN jobs j2 ON j1.company = j2.company
                AND LOWER(j1.title) = LOWER(j2.title)
                AND j1.job_id < j2.job_id
                AND ABS(EXTRACT(EPOCH FROM (j1.date_scraped - j2.date_scraped))/86400) <= 7
            WHERE j1.is_active=TRUE AND j2.is_active=TRUE
        ) sub
    """)
    return {"count": int(df["cnt"].iloc[0]) if not df.empty else 0}

@app.post("/api/jobs/deduplicate")
def deduplicate():
    """Mark older duplicates as inactive — keeps the newest posting."""
    conn = _conn(); cur = conn.cursor()
    cur.execute("""
        UPDATE jobs SET is_active=FALSE
        WHERE job_id IN (
            SELECT j1.job_id
            FROM jobs j1
            JOIN jobs j2 ON j1.company = j2.company
                AND LOWER(j1.title) = LOWER(j2.title)
                AND j1.date_scraped < j2.date_scraped
                AND ABS(EXTRACT(EPOCH FROM (j1.date_scraped - j2.date_scraped))/86400) <= 7
            WHERE j1.is_active=TRUE AND j2.is_active=TRUE
        )
    """)
    removed = cur.rowcount
    conn.commit(); conn.close()
    return {"status": "done", "removed": removed}


# ── Learning Path Resources ───────────────────────────────────────────────────

LEARNING_PATH = {
    "Python":       {"icon":"🐍","platform":"Real Python",       "url":"https://realpython.com"},
    "SQL":          {"icon":"🗄️","platform":"Mode SQL Tutorial",  "url":"https://mode.com/sql-tutorial"},
    "R":            {"icon":"📊","platform":"R for Data Science", "url":"https://r4ds.had.co.nz"},
    "Scala":        {"icon":"⚡","platform":"Rock the JVM",       "url":"https://rockthejvm.com"},
    "AWS":          {"icon":"☁️","platform":"AWS Skill Builder",  "url":"https://explore.skillbuilder.aws"},
    "Azure":        {"icon":"☁️","platform":"Microsoft Learn",    "url":"https://learn.microsoft.com/azure"},
    "GCP":          {"icon":"☁️","platform":"Google Cloud Skills","url":"https://cloudskillsboost.google"},
    "Snowflake":    {"icon":"❄️","platform":"Snowflake Uni",      "url":"https://learn.snowflake.com"},
    "Databricks":   {"icon":"⚡","platform":"Databricks Academy", "url":"https://www.databricks.com/learn"},
    "TensorFlow":   {"icon":"🧠","platform":"TensorFlow.org",     "url":"https://www.tensorflow.org/learn"},
    "PyTorch":      {"icon":"🔥","platform":"PyTorch Tutorials",  "url":"https://pytorch.org/tutorials"},
    "Scikit-learn": {"icon":"🤖","platform":"Scikit-learn Docs",  "url":"https://scikit-learn.org/stable/tutorial"},
    "Airflow":      {"icon":"🌊","platform":"Astronomer Academy", "url":"https://academy.astronomer.io"},
    "dbt":          {"icon":"🔧","platform":"dbt Learn",          "url":"https://courses.getdbt.com"},
    "Spark":        {"icon":"✨","platform":"Databricks Academy", "url":"https://www.databricks.com/learn"},
    "Kafka":        {"icon":"📨","platform":"Confluent Learn",    "url":"https://developer.confluent.io/learn"},
    "Docker":       {"icon":"🐳","platform":"Docker Docs",        "url":"https://docs.docker.com/get-started"},
    "Kubernetes":   {"icon":"⚙️","platform":"KodeKloud",          "url":"https://kodekloud.com/courses/kubernetes"},
    "Terraform":    {"icon":"🏗️","platform":"HashiCorp Learn",    "url":"https://developer.hashicorp.com/terraform/tutorials"},
    "Tableau":      {"icon":"📈","platform":"Tableau Learning",   "url":"https://www.tableau.com/learn/training"},
    "Power BI":     {"icon":"📊","platform":"Microsoft Learn",    "url":"https://learn.microsoft.com/power-bi"},
    "Machine Learning": {"icon":"🤖","platform":"fast.ai",       "url":"https://course.fast.ai"},
    "Deep Learning":    {"icon":"🧠","platform":"fast.ai",       "url":"https://course.fast.ai"},
    "LLM":          {"icon":"🤖","platform":"Hugging Face Course","url":"https://huggingface.co/learn/nlp-course"},
    "PostgreSQL":   {"icon":"🐘","platform":"PostgreSQL Tutorial","url":"https://www.postgresqltutorial.com"},
    "MongoDB":      {"icon":"🍃","platform":"MongoDB University", "url":"https://university.mongodb.com"},
    "Statistics":   {"icon":"📐","platform":"Khan Academy",       "url":"https://www.khanacademy.org/math/statistics-probability"},
    "MLOps":        {"icon":"⚙️","platform":"MLOps Community",   "url":"https://mlops.community"},
}

@app.get("/api/learning-path")
def learning_path(skills: str = Query(default="")):
    if not skills:
        return []
    skill_list = [s.strip() for s in skills.split(",") if s.strip()][:12]
    result = []
    for skill in skill_list:
        res = LEARNING_PATH.get(skill)
        if res:
            result.append({"skill": skill, **res})
        else:
            # Generic YouTube fallback
            result.append({
                "skill":    skill,
                "icon":     "📚",
                "platform": "YouTube",
                "url":      f"https://www.youtube.com/results?search_query={skill.replace(' ', '+')}+tutorial+beginners",
            })
    return result


# ── Interview Prep ────────────────────────────────────────────────────────────

@app.post("/api/resume/interview-prep")
async def interview_prep(request: Request):
    """Generate N likely interview Q&A pairs tailored to resume + job."""
    body    = await request.json()
    job_id  = body.get("job_id")
    resume_text = body.get("resume_text", "")[:3000]
    cfg     = _load_scrape_config()
    api_key = cfg.get("openai_api_key", "").strip()

    if not api_key:
        return JSONResponse({"error": "No OpenAI API key configured."}, status_code=400)

    job_info, job_desc = {}, ""
    if job_id:
        try:
            df = _q("SELECT title, company, location, description FROM jobs WHERE job_id=%s", [job_id])
            if not df.empty:
                row = df.iloc[0]
                job_info = {"title": str(row.get("title","") or ""), "company": str(row.get("company","") or "")}
                job_desc = str(row.get("description") or "")[:2000]
        except Exception:
            pass

    import openai
    client = openai.OpenAI(api_key=api_key)
    prompt = f"""Generate 8 tailored interview questions for this job application, with concise model answers based on the candidate's resume.

Job: {job_info.get('title','Unknown')} at {job_info.get('company','Unknown')}
Job Description: {job_desc or 'Not available'}
Candidate Resume: {resume_text}

Return JSON: {{"questions": [{{"q": "question", "a": "model answer using candidate's background"}}]}}"""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"user","content":prompt}],
        response_format={"type":"json_object"},
        temperature=0.5,
        max_tokens=2000,
    )
    return json.loads(resp.choices[0].message.content)


# ── Pipeline control ──────────────────────────────────────────────────────────

_running = threading.Event()

# ── Background scheduler ───────────────────────────────────────────────────────
# Tracks when the pipeline was last auto-run and fires it at the configured interval.

_sched_lock            = threading.Lock()
_sched_last_run: float = 0.0          # epoch seconds of last auto-run
_sched_stop            = threading.Event()


def _scheduler_loop():
    """Background daemon thread: fires run_full_pipeline() at the configured interval."""
    global _sched_last_run
    print("[scheduler] background thread started")
    while not _sched_stop.is_set():
        try:
            cfg = _load_scrape_config()
            enabled  = cfg.get("schedule_enabled", False)
            interval = float(cfg.get("schedule_interval_hours", 24))
        except Exception as e:
            print(f"[scheduler] config error: {e}")
            _sched_stop.wait(60)
            continue

        if enabled and interval > 0:
            with _sched_lock:
                elapsed_hours = (time.time() - _sched_last_run) / 3600.0
            if elapsed_hours >= interval:
                if not _running.is_set():
                    print(f"[scheduler] interval={interval}h elapsed, triggering pipeline …")
                    _running.set()
                    with _sched_lock:
                        _sched_last_run = time.time()
                    try:
                        from scheduler.runner import run_full_pipeline
                        run_full_pipeline()
                    except Exception as e:
                        print(f"[scheduler] pipeline error: {e}")
                    finally:
                        _running.clear()
                else:
                    print("[scheduler] pipeline already running, skipping tick")
        # Check every 60 seconds
        _sched_stop.wait(60)

    print("[scheduler] background thread stopped")


def _start_scheduler():
    t = threading.Thread(target=_scheduler_loop, name="auto-scheduler", daemon=True)
    t.start()


# Start the scheduler when the module loads (works with uvicorn + reload=False)
_start_scheduler()


@app.post("/api/pipeline/run")
def pipeline_run(bt: BackgroundTasks):
    if _running.is_set():
        return {"status": "already_running", "message": "Pipeline is already running"}

    def _go():
        _running.set()
        try:
            from scheduler.runner import run_full_pipeline
            run_full_pipeline()
        except Exception as e:
            print(f"[pipeline] error: {e}")
        finally:
            _running.clear()

    bt.add_task(_go)
    return {"status": "started", "message": "Pipeline started in background"}

@app.get("/api/pipeline/status")
def pipeline_status():
    cfg = _load_scrape_config()
    enabled  = cfg.get("schedule_enabled", False)
    interval = float(cfg.get("schedule_interval_hours", 24))
    with _sched_lock:
        last = _sched_last_run

    if last > 0:
        next_run_epoch = last + interval * 3600
        next_run_iso   = datetime.fromtimestamp(next_run_epoch).strftime("%Y-%m-%d %H:%M:%S")
        last_run_iso   = datetime.fromtimestamp(last).strftime("%Y-%m-%d %H:%M:%S")
        hours_until    = max(0.0, (next_run_epoch - time.time()) / 3600.0)
    else:
        next_run_iso = "Not yet scheduled" if enabled else "Scheduler disabled"
        last_run_iso = None
        hours_until  = interval if enabled else None

    return {
        "running":                _running.is_set(),
        "scheduler_enabled":      enabled,
        "schedule_interval_hours": interval,
        "last_auto_run":          last_run_iso,
        "next_auto_run":          next_run_iso,
        "hours_until_next_run":   round(hours_until, 2) if hours_until is not None else None,
    }

@app.get("/api/pipeline/logs")
def pipeline_logs(lines: int = Query(default=200, le=1000)):
    """Return the last N lines from today's log file."""
    log_dir  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
    log_file = os.path.join(log_dir, f"{date.today()}.log")
    if not os.path.exists(log_file):
        return {"lines": [], "file": log_file, "exists": False}
    try:
        with open(log_file, "r", encoding="utf-8", errors="replace") as f:
            all_lines = f.readlines()
        tail = [l.rstrip("\n") for l in all_lines[-lines:]]
        return {"lines": tail, "file": os.path.basename(log_file), "exists": True}
    except Exception as e:
        return {"lines": [f"Error reading log: {e}"], "file": log_file, "exists": False}

# ── Frontend ──────────────────────────────────────────────────────────────────

_here     = os.path.dirname(os.path.abspath(__file__))
_dash_dir = os.path.join(_here, "dashboard_new")
os.makedirs(_dash_dir, exist_ok=True)

@app.get("/")
def serve_root():
    f = os.path.join(_dash_dir, "index.html")
    return FileResponse(f) if os.path.exists(f) else JSONResponse({"error": "index.html not found"}, status_code=404)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api:app", host="0.0.0.0", port=8502, reload=False)
